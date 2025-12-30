import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
import time
import os
from dotenv import load_dotenv
from modules import news_collector, gemini_analyzer, github_storage

# Load environment variables - CRITICAL: override=True
load_dotenv(override=True)

# Page Config
st.set_page_config(
    page_title="Macoll Newsroom Dashboard",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .reportview-container {
        background: #f0f2f6
    }
    .metric-card {
        background-color: #ffffff;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .language-selector {
        position: fixed;
        top: 60px;
        right: 20px;
        z-index: 999;
        background: #000000;
        padding: 5px 10px;
        border-radius: 5px;
        font-size: 12px;
    }
    .language-selector button {
        background: transparent;
        color: white;
        border: none;
        padding: 5px 10px;
        cursor: pointer;
        font-size: 11px;
        font-weight: 600;
    }
    .language-selector button:hover {
        color: #3498db;
    }
    .language-selector button.active {
        color: #3498db;
        text-decoration: underline;
    }
    .stButton>button {
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

# Session State for Admin
if 'is_admin' not in st.session_state:
    st.session_state.is_admin = False

def run_analysis_pipeline(keyword, start_date, end_date):
    """Executes the full analysis pipeline."""
    placeholders = st.empty()
    with placeholders.container():
        st.info(f"🔍 Collecting news for '{keyword}' ({start_date} ~ {end_date})...")
        
        # 1. Collect News
        result = news_collector.search_naver_news(keyword, str(start_date), str(end_date))
        if not result['success']:
            st.error(f"Collection failed: {result.get('error')}")
            return False
            
        articles = result['article_details']
        total_count = len(articles)
        st.success(f"✅ Collected {total_count} articles.")
        
        if total_count == 0:
            st.warning("No articles found to analyze.")
            return False

        # 2. Analyze Sentiment
        st.info("🤖 Analyzing sentiment with Gemini...")
        try:
            sentiments = gemini_analyzer.analyze_sentiment_batch(articles)
        except Exception as e:
            st.error(f"❌ Sentiment analysis failed: {e}")
            st.warning("Using 'Neutral' sentiment for all articles as fallback.")
            sentiments = ["Neutral"] * len(articles)
        
        # Merge sentiment into articles
        for i, art in enumerate(articles):
            art['sentiment'] = sentiments[i] if i < len(sentiments) else "Neutral"
            
        # Calc Stats
        pos = sentiments.count('Positive')
        neg = sentiments.count('Negative')
        neu = sentiments.count('Neutral')
        
        # 3. Generate Report
        st.info("📝 Generating Issue Report...")
        sentiment_summary = f"Positive: {pos}, Negative: {neg}, Neutral: {neu}"
        try:
            report = gemini_analyzer.generate_issue_report(keyword, articles, sentiment_summary)
        except Exception as e:
            st.error(f"❌ Report generation failed: {e}")
            report = f"# Error\\n\\nFailed to generate report due to: {e}\\n\\nPlease check your GOOGLE_API_KEY and try again."
        
        # 4. Save Data
        data = {
            "keyword": keyword,
            "period": f"{start_date} ~ {end_date}",
            "summary_stats": { "positive": pos, "negative": neg, "neutral": neu },
            "report": report,
            "articles": articles,
            "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        st.info("💾 Saving to storage...")
        if github_storage.save_report(keyword, data):
            st.success("🎉 Analysis complete and saved!")
            time.sleep(2)
            placeholders.empty()
            return True
        else:
            st.error("Failed to save report.")
            return False

# Initialize language setting in session state
if 'language' not in st.session_state:
    st.session_state.language = 'Korean'

# Sidebar - Admin & Navigation
with st.sidebar:
    st.title("Macoll Newsroom")
    
    # Keyword Selection (Read Mode)
    try:
        available_keywords = github_storage.get_keyword_list()
    except Exception as e:
        st.error(f"Failed to load keywords: {e}")
        available_keywords = []
    
    selected_keyword = st.selectbox(
        "Select Issue Report", 
        options=["Select a keyword..."] + available_keywords,
        index=0
    )
    
    st.divider()
    
    # Admin Login / Panel
    if not st.session_state.is_admin:
        with st.expander("🔐 Admin Access"):
            password = st.text_input("Password", type="password")
            if st.button("Login"):
                admin_password = os.getenv("ADMIN_PASSWORD")
                if not admin_password:
                    st.error("⚠️ Admin password not configured. Please set ADMIN_PASSWORD in .env file.")
                elif password == admin_password:
                    st.session_state.is_admin = True
                    st.rerun()
                else:
                    st.error("Invalid password")
    else:
        st.subheader("Admin Panel")
        if st.button("Logout"):
            st.session_state.is_admin = False
            st.rerun()
        
        # Keyword Management Section
        st.divider()
        st.markdown("### Saved Keywords")
        
        try:
            saved_keywords = github_storage.get_keyword_list()
            
            if saved_keywords:
                st.caption(f"Total: {len(saved_keywords)} keywords")
                
                # Display keywords with delete buttons
                for keyword in sorted(saved_keywords):
                    col1, col2 = st.columns([4, 1])
                    with col1:
                        st.write(f"**{keyword}**")
                    with col2:
                        if st.button("❌", key=f"delete_{keyword}"):
                            if github_storage.delete_report(keyword):
                                st.success(f"Deleted: {keyword}")
                                st.rerun()
                            else:
                                st.error(f"Failed to delete: {keyword}")
            else:
                st.info("No saved keywords found.")
        except Exception as e:
            st.error(f"Error loading keywords: {e}")
        
        st.divider()
            
        st.markdown("### New Analysis")
        new_keyword = st.text_input("Target Keyword")
        
        col1, col2 = st.columns(2)
        with col1:
            start_d = st.date_input("Start Date", datetime.now())
        with col2:
            end_d = st.date_input("End Date", datetime.now())
            
        if st.button("Run Analysis", type="primary"):
            if new_keyword:
                with st.spinner("Processing..."):
                    if run_analysis_pipeline(new_keyword, start_d, end_d):
                        st.balloons()
                        st.success(f"✅ '{new_keyword}' analysis complete! Refreshing...")
                        time.sleep(2)  # Show success message before refresh
                        st.rerun()  # Refresh to update keyword list
            else:
                st.warning("Please enter a keyword.")

# Main Content
if selected_keyword and selected_keyword != "Select a keyword...":
    data = github_storage.load_report(selected_keyword)
    
    if data:
        # Language Selector (Top Right) - Custom Styled
        st.markdown("""
        <style>
        /* Language buttons - force black styling with maximum specificity */
        button[data-testid="baseButton-secondary"],
        button[data-testid="baseButton-primary"] {
            font-size: 10px !important;
            padding: 3px 6px !important;
            min-height: 24px !important;
            background-color: #ffffff !important;
            color: #000000 !important;
            border: 1px solid #e0e0e0 !important;
            font-weight: 400 !important;
        }
        button[data-testid="baseButton-primary"],
        button[data-testid="baseButton-primary"]:hover,
        button[data-testid="baseButton-primary"]:active,
        button[data-testid="baseButton-primary"]:focus {
            background-color: #000000 !important;
            background-image: none !important;
            color: white !important;
            border: 1px solid #000000 !important;
            font-weight: 600 !important;
        }
        </style>
        """, unsafe_allow_html=True)
               
        # 제목과 언어 선택 버튼을 같은 줄에 배치
        header_col1, header_col2 = st.columns([3, 1])
        
        with header_col1:
            st.markdown(f"## Issue Report: {data['keyword']}")
        
        with header_col2:
            # 언어 선택 버튼 (우측 정렬)
            lang_col1, lang_col2 = st.columns(2)
            with lang_col1:
                if st.button("KR", key="lang_kr", use_container_width=True, 
                           type="primary" if st.session_state.get('report_lang', 'KR') == 'KR' else "secondary"):
                    st.session_state.report_lang = 'KR'
                    st.rerun()
            with lang_col2:
                if st.button("EN", key="lang_en", use_container_width=True,
                           type="primary" if st.session_state.get('report_lang', 'KR') == 'EN' else "secondary"):
                    st.session_state.report_lang = 'EN'
                    st.rerun()
        
        # 기간 정보
        st.caption(f"Period: {data['period']} | Last Updated: {data.get('updated_at', '-')}")
        
        # Summary Stats Cards
        stats = data['summary_stats']
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.metric("Total Articles", len(data['articles']))
        with c2:
            st.metric("Positive", stats['positive'], delta_color="normal")
        with c3:
            st.metric("Negative", stats['negative'], delta_color="inverse")
        with c4:
            st.metric("Neutral", stats['neutral'], delta_color="off")
            
        st.divider()
        
        # Content Layout
        col_left, col_right = st.columns([1, 1], gap="large")
        
        with col_left:
            st.subheader("Sentiment Timeline")
            
            # Prepare Data for Charts
            df = pd.DataFrame(data['articles'])
            
            # Initialize session state for selected date
            if 'selected_date' not in st.session_state:
                st.session_state.selected_date = None
            
            # Ensure date format
            if not df.empty and 'date' in df.columns:
                # Try to parse dates - handle both datetime and string formats
                def parse_date(date_val):
                    if pd.isna(date_val):
                        return None
                    # If already a string in YYYY-MM-DD format, keep it
                    if isinstance(date_val, str):
                        # Check if it's already in YYYY-MM-DD format
                        if len(date_val) == 10 and date_val[4] == '-' and date_val[7] == '-':
                            return date_val
                        # Try to parse other string formats
                        try:
                            return pd.to_datetime(date_val).strftime('%Y-%m-%d')
                        except:
                            return None
                    # If it's a datetime object
                    try:
                        return pd.to_datetime(date_val).strftime('%Y-%m-%d')
                    except:
                        return None
                
                df['date'] = df['date'].apply(parse_date)
                df = df.dropna(subset=['date'])  # Remove rows with invalid dates
                
                # Get date range from data period
                period_str = data.get('period', '')
                if '~' in period_str:
                    try:
                        start_str, end_str = period_str.split('~')
                        start_date = pd.to_datetime(start_str.strip()).strftime('%Y-%m-%d')
                        end_date = pd.to_datetime(end_str.strip()).strftime('%Y-%m-%d')
                    except:
                        # Fallback to min/max from data
                        start_date = df['date'].min()
                        end_date = df['date'].max()
                else:
                    start_date = df['date'].min()
                    end_date = df['date'].max()
                
                # Create complete date range
                all_dates = pd.date_range(start=start_date, end=end_date, freq='D')
                all_dates_str = [d.strftime('%Y-%m-%d') for d in all_dates]
                
                # Aggregation by Date and Sentiment
                daily_sentiment = df.groupby(['date', 'sentiment']).size().reset_index(name='count')
                daily_total = df.groupby('date').size().reset_index(name='total')
                
                # Debug: Print sentiment distribution
                print("=== DEBUG: Sentiment Distribution ===")
                print(daily_sentiment.head(20))
                print(f"\nTotal Positive: {df[df['sentiment'] == 'Positive'].shape[0]}")
                print(f"Total Negative: {df[df['sentiment'] == 'Negative'].shape[0]}")
                print(f"Total Neutral: {df[df['sentiment'] == 'Neutral'].shape[0]}")
                
                # Create complete dataframes with all dates for each sentiment
                sentiments = ['Positive', 'Negative', 'Neutral']
                complete_sentiment_data = {}
                
                for sentiment in sentiments:
                    sentiment_df = daily_sentiment[daily_sentiment['sentiment'] == sentiment]
                    # Create complete date range dataframe
                    complete_df = pd.DataFrame({'date': all_dates_str})
                    # Merge with actual data
                    complete_df = complete_df.merge(sentiment_df[['date', 'count']], on='date', how='left')
                    # Fill missing values with 0
                    complete_df['count'] = complete_df['count'].fillna(0).astype(int)
                    complete_sentiment_data[sentiment] = complete_df
                
                # Create complete total dataframe
                complete_total = pd.DataFrame({'date': all_dates_str})
                complete_total = complete_total.merge(daily_total, on='date', how='left')
                complete_total['total'] = complete_total['total'].fillna(0).astype(int)
                
                # Overlaid Area Chart: Each sentiment fills from zero independently
                fig = go.Figure()
                
                # Add areas in order from largest to smallest for better visibility
                # Add Negative area (red) - usually largest, so add first
                negative_data = complete_sentiment_data['Negative']
                fig.add_trace(go.Scatter(
                    x=negative_data['date'].tolist(),
                    y=negative_data['count'].tolist(),
                    mode='lines',
                    name='부정 (Negative)',
                    line=dict(color='#e74c3c', width=2, shape='spline'),  # Increased width for visibility
                    fill='tozeroy',
                    fillcolor='rgba(231, 76, 60, 0.5)',
                    hovertemplate='<b>부정</b><br>%{y}<extra></extra>'
                ))
                
                # Add Positive area (blue)
                positive_data = complete_sentiment_data['Positive']
                fig.add_trace(go.Scatter(
                    x=positive_data['date'].tolist(),
                    y=positive_data['count'].tolist(),
                    mode='lines',
                    name='긍정 (Positive)',
                    line=dict(color='#3498db', width=2, shape='spline'),  # Increased width for visibility
                    fill='tozeroy',
                    fillcolor='rgba(52, 152, 219, 0.5)',
                    hovertemplate='<b>긍정</b><br>%{y}<extra></extra>'
                ))
                
                # Add Neutral area (gray) - smallest, add last so it's on top
                neutral_data = complete_sentiment_data['Neutral']
                fig.add_trace(go.Scatter(
                    x=neutral_data['date'].tolist(),
                    y=neutral_data['count'].tolist(),
                    mode='lines',
                    name='중립 (Neutral)',
                    line=dict(color='#95a5a6', width=2, shape='spline'),  # Increased width for visibility
                    fill='tozeroy',
                    fillcolor='rgba(149, 165, 166, 0.5)',
                    hovertemplate='<b>중립</b><br>%{y}<extra></extra>'
                ))
                
                # Add Total line (black dashed) - separate from stack
                fig.add_trace(go.Scatter(
                    x=complete_total['date'].tolist(),
                    y=complete_total['total'].tolist(),
                    mode='lines+markers',
                    name='전체 기사량',
                    line=dict(color='#000000', width=2, dash='dot', shape='spline'),
                    marker=dict(size=6, color='#000000'),
                    hovertemplate='<b>전체</b><br>%{y}<extra></extra>'
                ))
                
                fig.update_layout(
                    title=None,  # Removed title as requested
                    xaxis_title=None,
                    yaxis_title=None,  # Removed Y-axis title as requested
                    hovermode='x unified',
                    height=450,
                    margin=dict(t=40, b=80, l=50, r=20),  # Increased bottom margin for date labels
                    xaxis=dict(
                        type='category',
                        tickangle=-45,
                        showgrid=True,
                        gridcolor='rgba(128, 128, 128, 0.2)'
                    ),
                    yaxis=dict(
                        showgrid=True,
                        gridcolor='rgba(128, 128, 128, 0.2)'
                    ),
                    legend=dict(
                        orientation="h",
                        yanchor="top",
                        y=1.02,  # Adjusted legend position
                        xanchor="right",
                        x=1
                    ),
                    plot_bgcolor='white',
                    paper_bgcolor='white'
                )

                # Use streamlit-plotly-events to capture click events
                from streamlit_plotly_events import plotly_events
                
                # Display chart with click event handling
                selected_points = plotly_events(fig, click_event=True, hover_event=False, select_event=False, override_height=450, override_width="100%", key="sentiment_chart")
                
                # Handle click event
                if selected_points and len(selected_points) > 0:
                    clicked_point = selected_points[0]
                    if 'x' in clicked_point:
                        clicked_date = clicked_point['x']
                        # Only update if it's a different date
                        if st.session_state.get('selected_date') != clicked_date:
                            st.session_state.selected_date = clicked_date
                            st.rerun()
                
                # Article List
                st.subheader("Recent Articles")
                
                # Show "Show All Dates" button AFTER "Recent Articles" when a date is selected
                if st.session_state.selected_date:
                    st.info(f"📅 Filtered by date: **{st.session_state.selected_date}**")
                    if st.button("🔄 Show All Dates", key="show_all_dates_btn", use_container_width=True, type="primary"):
                        st.session_state.selected_date = None
                        st.session_state.sentiment_filter = ['Positive', 'Negative', 'Neutral']
                        st.rerun()
                
                # Initialize sentiment filter in session state if not exists
                if 'sentiment_filter' not in st.session_state:
                    st.session_state.sentiment_filter = ['Positive', 'Negative', 'Neutral']
                
                # Date filter based on selected date
                if st.session_state.selected_date:
                    filtered_df = df[df['date'] == st.session_state.selected_date]
                else:
                    # Sentiment filter
                    sentiment_filter = st.multiselect(
                        "Filter by Sentiment", 
                        ['Positive', 'Negative', 'Neutral'], 
                        default=st.session_state.sentiment_filter,
                        key='sentiment_multiselect'
                    )
                    # Update session state
                    st.session_state.sentiment_filter = sentiment_filter
                    filtered_df = df[df['sentiment'].isin(sentiment_filter)]
                
                # Limit to 30 articles and sort by date (newest first)
                filtered_df = filtered_df.sort_values('date', ascending=False).head(30)
                
                # Display count
                st.caption(f"Showing {len(filtered_df)} articles (max 30)")
                
                # CSS for article cards
                st.markdown("""
                <style>
                .article-card {
                    border-left: 4px solid;
                    padding: 8px 12px;
                    margin-bottom: 8px;
                    border-radius: 4px;
                    transition: transform 0.2s;
                }
                .article-card:hover {
                    transform: translateX(4px);
                }
                .article-card.positive {
                    background-color: #e8f4fd;
                    border-left-color: #3498db;
                }
                .article-card.negative {
                    background-color: #fde8e8;
                    border-left-color: #e74c3c;
                }
                .article-card.neutral {
                    background-color: #f5f5f5;
                    border-left-color: #95a5a6;
                }
                .article-row {
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    width: 100%;
                }
                .article-press {
                    font-size: 11px;
                    color: #2c3e50;
                    font-weight: 700;
                }
                .article-date {
                    font-size: 11px;
                    color: #95a5a6;
                    font-weight: 500;
                }
                .article-title {
                    font-size: 14px;
                    color: #34495e;
                    font-weight: 400;
                    margin-top: 4px;
                    flex: 1;
                    padding-right: 10px;
                }
                .article-title a {
                    color: #34495e;
                    text-decoration: none;
                }
                .article-title a:hover {
                    text-decoration: underline;
                }
                .sentiment-badge {
                    font-size: 10px;
                    padding: 2px 8px;
                    border-radius: 3px;
                    font-weight: 600;
                    white-space: nowrap;
                }
                .sentiment-badge.positive {
                    background-color: #3498db;
                    color: white;
                }
                .sentiment-badge.negative {
                    background-color: #e74c3c;
                    color: white;
                }
                .sentiment-badge.neutral {
                    background-color: #95a5a6;
                    color: white;
                }
                </style>
                """, unsafe_allow_html=True)
                
                # Display article cards
                for _, row in filtered_df.iterrows():
                    sentiment_class = row['sentiment'].lower()
                    sentiment_kr = {'positive': '긍정', 'negative': '부정', 'neutral': '중립'}
                    sentiment_label = sentiment_kr.get(sentiment_class, sentiment_class)
                    
                    st.markdown(f"""
                    <div class="article-card {sentiment_class}">
                        <!-- Row 1: Press and Date -->
                        <div class="article-row">
                            <div class="article-press">{row['press']}</div>
                            <div class="article-date">{row['date']}</div>
                        </div>
                        <!-- Row 2: Title and Sentiment -->
                        <div class="article-row" style="margin-top: 4px; align-items: flex-end;">
                            <div class="article-title">
                                <a href="{row['link']}" target="_blank">{row['title']}</a>
                            </div>
                            <div class="sentiment-badge {sentiment_class}">{sentiment_label}</div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Excel Download (Moved to left column bottom)
                if not df.empty:
                    try:
                        from io import BytesIO
                        from datetime import datetime
                        
                        # Reorder columns: date, press, title, link, sentiment
                        excel_df = df[['date', 'press', 'title', 'link', 'sentiment']].copy()
                        
                        # Standardize date for excel
                        if 'date' in excel_df.columns:
                            excel_df['date'] = excel_df['date'].astype(str)
                        
                        output = BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            excel_df.to_excel(writer, index=False, sheet_name='News Analysis')
                        output.seek(0)
                        
                        st.download_button(
                            label="Download as Excel",
                            data=output,
                            file_name=f"news_analysis_{data['keyword']}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    except ImportError:
                        st.info("💡 Install openpyxl to enable Excel export")
            else:
                st.info("No data available for visualization.")

        with col_right:
            st.subheader("Issue Analysis Report")
            
            report_content = data.get('report', "No report content.")
            
            # Initialize translation cache in session state
            if 'translated_reports' not in st.session_state:
                st.session_state.translated_reports = {}
            
            # Create cache key based on keyword and original report hash
            cache_key = f"{data['keyword']}_{hash(report_content)}"
            
            # Auto-translate if language is English
            if st.session_state.language == 'English' and report_content != "No report content.":
                # Check if translation is already cached
                if cache_key in st.session_state.translated_reports:
                    report_content = st.session_state.translated_reports[cache_key]
                    st.info("📝 Showing cached English translation")
                else:
                    with st.spinner("🌐 Translating report to English..."):
                        try:
                            translated = gemini_analyzer.translate_report(report_content, "English")
                            st.session_state.translated_reports[cache_key] = translated
                            report_content = translated
                            st.success("✅ Translation complete!")
                        except Exception as e:
                            st.warning(f"⚠️ Translation failed: {e}. Showing original Korean report.")
            
            
            # Remove the first H1 heading (main report title)
            import re
            if report_content != "No report content.":
                # Remove the first H1 heading (# Title)
                report_content = re.sub(r'^#\s+[^\n]+\n+', '', report_content, count=1, flags=re.MULTILINE)
            
            # Add CSS for consistent line spacing in the report
            st.markdown("""
                <style>
                /* Ensure consistent line spacing for all report content */
                .stMarkdown p {
                    line-height: 1.6 !important;
                    margin-bottom: 1em !important;
                }
                .stMarkdown ul, .stMarkdown ol {
                    line-height: 1.6 !important;
                    margin-bottom: 1em !important;
                }
                .stMarkdown li {
                    line-height: 1.6 !important;
                    margin-bottom: 0.5em !important;
                }
                .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4 {
                    line-height: 1.4 !important;
                    margin-top: 0.5em !important;
                    margin-bottom: 0.4em !important;
                }
                .stMarkdown table {
                    line-height: 1.6 !important;
                }
                </style>
            """, unsafe_allow_html=True)
            
            st.markdown(report_content, unsafe_allow_html=True)
            
            # Word Download Button - Comprehensive Export
            if report_content != "No report content.":
                try:
                    from io import BytesIO
                    from docx import Document
                    from docx.shared import Pt, RGBColor, Inches
                    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    import re
                    import tempfile
                    import os as os_module
                    
                    # Create Word document
                    doc = Document()
                    
                    # Set default font to 맑은고딕 (Malgun Gothic)
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = '맑은고딕'
                    font.size = Pt(10)
                    
                    # For compatibility with older Word versions
                    rFonts = style.element.rPr.rFonts
                    rFonts.set(qn('w:eastAsia'), '맑은고딕')
                    
                    # Add Title
                    title = doc.add_heading(f"Issue Report: {data['keyword']}", level=0)
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Add Period Info
                    period_para = doc.add_paragraph()
                    period_para.add_run(f"Period: {data['period']} | Last Updated: {data.get('updated_at', '-')}").italic = True
                    period_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.add_paragraph()  # Spacing
                    
                    # Add Summary Statistics
                    doc.add_heading('Summary Statistics', level=1)
                    stats_table = doc.add_table(rows=1, cols=4)
                    stats_table.style = 'Light Grid Accent 1'
                    hdr_cells = stats_table.rows[0].cells
                    hdr_cells[0].text = 'Total Articles'
                    hdr_cells[1].text = 'Positive'
                    hdr_cells[2].text = 'Negative'
                    hdr_cells[3].text = 'Neutral'
                    
                    row_cells = stats_table.add_row().cells
                    row_cells[0].text = str(len(data['articles']))
                    row_cells[1].text = str(stats['positive'])
                    row_cells[2].text = str(stats['negative'])
                    row_cells[3].text = str(stats['neutral'])
                    doc.add_paragraph()  # Spacing
                    
                    # Add Sentiment Timeline Chart as Image
                    if not df.empty:
                        doc.add_heading('Sentiment Timeline', level=1)
                        try:
                            # Recreate the chart (same as displayed)
                            daily_sentiment = df.groupby(['date', 'sentiment']).size().reset_index(name='count')
                            daily_total = df.groupby('date').size().reset_index(name='total')
                            
                            fig = go.Figure()
                            
                            # Add traces
                            positive_data = daily_sentiment[daily_sentiment['sentiment'] == 'Positive']
                            if not positive_data.empty:
                                fig.add_trace(go.Scatter(
                                    x=positive_data['date'], y=positive_data['count'],
                                    mode='lines', name='긍정 (Positive)',
                                    line=dict(color='#3498db', width=2, shape='spline'),
                                    fill='tozeroy', fillcolor='rgba(52, 152, 219, 0.3)'
                                ))
                            
                            negative_data = daily_sentiment[daily_sentiment['sentiment'] == 'Negative']
                            if not negative_data.empty:
                                fig.add_trace(go.Scatter(
                                    x=negative_data['date'], y=negative_data['count'],
                                    mode='lines', name='부정 (Negative)',
                                    line=dict(color='#e74c3c', width=2, shape='spline'),
                                    fill='tozeroy', fillcolor='rgba(231, 76, 60, 0.3)'
                                ))
                            
                            neutral_data = daily_sentiment[daily_sentiment['sentiment'] == 'Neutral']
                            if not neutral_data.empty:
                                fig.add_trace(go.Scatter(
                                    x=neutral_data['date'], y=neutral_data['count'],
                                    mode='lines', name='중립 (Neutral)',
                                    line=dict(color='#95a5a6', width=2, shape='spline'),
                                    fill='tozeroy', fillcolor='rgba(149, 165, 166, 0.3)'
                                ))
                            
                            fig.add_trace(go.Scatter(
                                x=daily_total['date'], y=daily_total['total'],
                                mode='lines+markers', name='전체 기사량',
                                line=dict(color='#000000', width=2, dash='dot', shape='spline'),
                                marker=dict(size=6, color='#000000')
                            ))
                            
                            fig.update_layout(
                                title="SENTIMENT TREND",
                                height=400, width=800,
                                xaxis=dict(type='category', tickangle=-45),
                                legend=dict(orientation="h", yanchor="top", y=1.15, xanchor="right", x=1)
                            )
                            
                            # Save chart as image
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                                chart_path = tmp.name
                                fig.write_image(chart_path, width=1200, height=600)
                            
                            # Add image to document
                            doc.add_picture(chart_path, width=Inches(6))
                            os_module.remove(chart_path)  # Clean up temp file
                            doc.add_paragraph()  # Spacing
                        except Exception as chart_error:
                            doc.add_paragraph(f"Chart generation failed: {chart_error}")
                    
                    # Add Article List Table
                    doc.add_heading('Recent Articles (Top 30)', level=1)
                    if not df.empty:
                        article_table = doc.add_table(rows=1, cols=4)
                        article_table.style = 'Light List Accent 1'
                        hdr_cells = article_table.rows[0].cells
                        hdr_cells[0].text = 'Date'
                        hdr_cells[1].text = 'Press'
                        hdr_cells[2].text = 'Title'
                        hdr_cells[3].text = 'Sentiment'
                        
                        # Add up to 30 articles
                        for _, row in df.head(30).iterrows():
                            row_cells = article_table.add_row().cells
                            row_cells[0].text = str(row['date'])
                            row_cells[1].text = str(row['press'])
                            row_cells[2].text = str(row['title'])
                            row_cells[3].text = str(row['sentiment'])
                        doc.add_paragraph()  # Spacing
                    
                    # Add Analysis Report
                    doc.add_heading('Detailed Analysis Report', level=1)
                    
                    # Parse markdown report and add to Word
                    lines = report_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Remove HTML tags like <mark>
                        line = re.sub(r'<[^>]+>', '', line)
                        
                        # Remove markdown bold markers (**)
                        line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
                        
                        # Headers
                        if line.startswith('# '):
                            p = doc.add_heading(line[2:], level=1)
                        elif line.startswith('## '):
                            p = doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            p = doc.add_heading(line[4:], level=3)
                        elif line.startswith('#### '):
                            p = doc.add_heading(line[5:], level=4)
                        # List items
                        elif line.startswith('- ') or line.startswith('* '):
                            p = doc.add_paragraph(line[2:], style='List Bullet')
                        elif re.match(r'^\d+\.', line):
                            p = doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
                        else:
                            p = doc.add_paragraph(line)
                    
                    # Save to buffer
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        label="Download as Word",
                        data=buffer,
                        file_name=f"{data['keyword']}_comprehensive_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except ImportError as ie:
                    st.info(f"💡 Install missing packages: {ie}")
                except Exception as e:
                    st.warning(f"Word export failed: {e}")

    else:
        st.error(f"Failed to load data for {selected_keyword}")
else:
    # Landing Page
    st.markdown("""
    ## Welcome to Macoll Newsroom
    
    This dashboard provides AI-powered insights into news trends.
    
    **Select a keyword** from the sidebar to view an existing report.
    
    **Login as Admin** in the sidebar to create new analysis reports.
    """)
